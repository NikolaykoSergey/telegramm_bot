import logging
from pathlib import Path
from typing import List, Dict

import PyPDF2
import docx
# импорт для OCR / таблиц, если используются
# import cv2
# import pytesseract

logger = logging.getLogger(__name__)

MAX_CHUNK_SIZE = 800  # символов в одном фрагменте


def split_text_into_chunks(text: str, max_size: int = MAX_CHUNK_SIZE) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    return [text[i:i + max_size] for i in range(0, len(text), max_size)]


class DocumentProcessor:
    """Обработка документов (PDF, DOCX, DOC) в список фрагментов для индексации"""

    def __init__(self):
        logger.info("✅ DocumentProcessor инициализирован")

    def process_file(self, file_path: Path) -> List[Dict]:
        logger.info("📄 Обработка файла: %s", file_path.name)
        suffix = file_path.suffix.lower()

        try:
            if suffix == ".pdf":
                return self._process_pdf(file_path)
            elif suffix in (".docx", ".doc"):
                return self._process_docx(file_path)
            else:
                logger.warning("⚠️ Неподдерживаемый тип файла: %s", suffix)
                return []
        except Exception as e:
            logger.error("❌ Ошибка при обработке файла %s: %s", file_path, repr(e))
            return []

    def _process_pdf(self, file_path: Path) -> List[Dict]:
        fragments: List[Dict] = []

        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            logger.info("📄 PDF: %s, страниц: %d", file_path.name, num_pages)

            for page_idx in range(num_pages):
                try:
                    page = reader.pages[page_idx]
                    text = page.extract_text() or ""
                except Exception as e:
                    logger.error("❌ Ошибка чтения страницы %d: %s", page_idx + 1, repr(e))
                    text = ""

                # режем текст на чанки
                for chunk in split_text_into_chunks(text):
                    fragments.append({
                        "content": chunk,
                        "page": page_idx + 1,
                        "type": "text",
                    })

                # тут же можно добавить OCR/таблицы, если у тебя они есть

        logger.info("📄 PDF %s: всего фрагментов %d", file_path.name, len(fragments))
        return fragments

    def _process_docx(self, file_path: Path) -> List[Dict]:
        fragments: List[Dict] = []

        doc = docx.Document(str(file_path))
        logger.info("📄 DOCX: %s, параграфов: %d", file_path.name, len(doc.paragraphs))

        page = 1  # если реальных страниц нет, можно считать всё страницей 1

        for para in doc.paragraphs:
            text = para.text or ""
            for chunk in split_text_into_chunks(text):
                fragments.append({
                    "content": chunk,
                    "page": page,
                    "type": "text",
                })

        logger.info("📄 DOCX %s: всего фрагментов %d", file_path.name, len(fragments))
        return fragments